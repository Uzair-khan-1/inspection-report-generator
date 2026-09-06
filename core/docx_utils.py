"""
DOCX handling for the AI Resume Tailor app.

Three jobs live here:
  1. extract_resume_text()   -> pull plain text out of an uploaded master resume
  2. detect_template_sections() -> quick read-only peek at a template's section headings
  3. generate_tailored_docx() -> the real work: populate the uploaded template with
     AI-tailored content while reusing the template's own paragraph/run formatting,
     so fonts, bullet styles, spacing, and tables are preserved.

Strategy for (3): rather than building a document from scratch (which would lose the
template's look), we treat existing paragraphs in the template as "formatting
exemplars". We clone the underlying XML of an exemplar paragraph for every new line/
bullet we need to add, set its text, and splice it back into the document next to
where the original content lived. This keeps fonts, bullet numbering, spacing, and
table structures intact even when the tailored content has a different number of
bullets/jobs than the template's placeholder example.

This is a best-effort approach and works well for standard single-column ATS resume
templates (heading + paragraphs/bullets). Extremely exotic templates (text boxes,
multi-column tables used for layout, SmartArt) may not be preserved perfectly.
"""

from __future__ import annotations

import copy
import io
from dataclasses import dataclass, field

from docx import Document
from docx.text.paragraph import Paragraph

# ---------------------------------------------------------------------------
# Section vocabulary
# ---------------------------------------------------------------------------

CANONICAL_ORDER = [
    "summary",
    "skills",
    "experience",
    "education",
    "certifications",
    "training",
    "projects",
    "languages",
]

SECTION_KEYWORDS = {
    "summary": [
        "professional summary", "career summary", "summary", "profile",
        "objective", "about me", "career objective",
    ],
    "skills": [
        "core competencies", "technical skills", "key skills", "skills",
        "competencies", "areas of expertise",
    ],
    "experience": [
        "professional experience", "work experience", "employment history",
        "career history", "experience",
    ],
    "education": ["education", "academic background", "academic qualifications"],
    "certifications": ["certifications & registrations", "certifications", "certificates", "licenses", "licensure", "registrations"],
    "training": ["professional training", "trainings", "training & development", "training and development", "training"],
    "projects": ["projects", "key projects", "selected projects"],
    "languages": ["languages", "language proficiency"],
}


def match_section(heading_text: str) -> str | None:
    text = heading_text.strip().lower()
    if not text:
        return None
    best = None
    best_len = 0
    for canonical, keywords in SECTION_KEYWORDS.items():
        for kw in keywords:
            if kw in text and len(kw) > best_len:
                best = canonical
                best_len = len(kw)
    return best


# ---------------------------------------------------------------------------
# Low level paragraph helpers
# ---------------------------------------------------------------------------

def _has_numpr(paragraph: Paragraph) -> bool:
    """True if this paragraph is a bulleted/numbered list item, whether the
    numbering is set directly on the paragraph or inherited from its style
    (most Word "List Bullet" / "List Number" styles work the latter way)."""
    pPr = paragraph._p.pPr
    if pPr is not None and pPr.numPr is not None:
        return True
    style_name = (paragraph.style.name or "") if paragraph.style else ""
    return any(kw in style_name for kw in ("List", "Bullet", "Number"))


def _is_heading(paragraph: Paragraph) -> bool:
    """A paragraph counts as a section heading only if it's a real Word
    heading/title style, OR it's short + visually emphasized (bold/all-caps/
    underlined) AND its text actually matches a known resume section keyword.
    The keyword requirement is what stops short bold lines like a person's
    name or a job title from being mistaken for a section heading."""
    text = paragraph.text.strip()
    if not text or _has_numpr(paragraph):
        return False
    style_name = (paragraph.style.name or "").lower() if paragraph.style else ""
    if "heading" in style_name or "title" in style_name:
        return True
    words = text.split()
    if len(words) <= 6 and match_section(text) is not None:
        bold = any(r.bold for r in paragraph.runs if r.bold)
        allcaps = text.upper() == text and any(c.isalpha() for c in text)
        underline = any(r.underline for r in paragraph.runs if r.underline)
        if bold or allcaps or underline:
            return True
    return False


def set_para_text(paragraph: Paragraph, text: str) -> None:
    """Set a paragraph's visible text while reusing the formatting of its
    first run (font, size, color, bold, etc.). Extra runs are emptied."""
    runs = paragraph.runs
    if runs:
        runs[0].text = text
        for r in runs[1:]:
            r.text = ""
    else:
        paragraph.add_run(text)


def _clone_paragraph_element(exemplar: Paragraph, text: str):
    """Deep-copy an exemplar paragraph's XML, set new text, return the raw
    <w:p> element (not yet attached to any tree)."""
    new_p = copy.deepcopy(exemplar._p)
    new_para = Paragraph(new_p, exemplar._parent)
    set_para_text(new_para, text)
    return new_p


def _remove_paragraph(paragraph: Paragraph) -> None:
    el = paragraph._p
    parent = el.getparent()
    if parent is not None:
        parent.remove(el)


def _insert_after(anchor_element, new_element):
    anchor_element.addnext(new_element)
    return new_element


# ---------------------------------------------------------------------------
# Template structure analysis
# ---------------------------------------------------------------------------

@dataclass
class SectionInfo:
    canonical: str | None
    heading_text: str
    heading_para: Paragraph
    content_paras: list = field(default_factory=list)

    def exemplars(self):
        """Return (header_lines[:2], bullet_exemplar, generic_exemplar)."""
        header_lines = []
        bullet = None
        generic = None
        seen_bullet = False
        for p in self.content_paras:
            if not p.text.strip():
                continue
            if generic is None:
                generic = p
            if _has_numpr(p):
                seen_bullet = True
                if bullet is None:
                    bullet = p
            else:
                if not seen_bullet and len(header_lines) < 2:
                    header_lines.append(p)
        return header_lines, bullet, generic


@dataclass
class TemplateProfile:
    document: Document
    contact_paras: list
    sections: list  # list[SectionInfo] in document order
    fallback_heading: Paragraph | None
    fallback_bullet: Paragraph | None
    fallback_plain: Paragraph | None

    def get(self, canonical: str) -> SectionInfo | None:
        for s in self.sections:
            if s.canonical == canonical:
                return s
        return None

    def detected_names(self) -> list[str]:
        return [s.heading_text for s in self.sections]


def _build_profile(document: Document) -> TemplateProfile:
    paragraphs = document.paragraphs
    heading_idx = []
    for idx, p in enumerate(paragraphs):
        if _is_heading(p):
            heading_idx.append(idx)

    contact_paras = paragraphs[: heading_idx[0]] if heading_idx else list(paragraphs)

    sections = []
    for i, idx in enumerate(heading_idx):
        start = idx + 1
        end = heading_idx[i + 1] if i + 1 < len(heading_idx) else len(paragraphs)
        content = paragraphs[start:end]
        canonical = match_section(paragraphs[idx].text)
        sections.append(SectionInfo(canonical, paragraphs[idx].text.strip(), paragraphs[idx], content))

    fallback_heading = paragraphs[heading_idx[0]] if heading_idx else None
    fallback_bullet = None
    fallback_plain = None
    for p in paragraphs:
        if p.text.strip():
            if _has_numpr(p) and fallback_bullet is None:
                fallback_bullet = p
            if not _has_numpr(p) and fallback_plain is None and p not in contact_paras:
                fallback_plain = p
        if fallback_bullet and fallback_plain:
            break

    return TemplateProfile(document, contact_paras, sections, fallback_heading, fallback_bullet, fallback_plain)


def detect_template_sections(template_bytes: bytes) -> list[str]:
    """Cheap read-only peek used by the UI to show what the app found in the
    uploaded template before generation happens."""
    document = Document(io.BytesIO(template_bytes))
    profile = _build_profile(document)
    return profile.detected_names()


# ---------------------------------------------------------------------------
# Master resume text extraction
# ---------------------------------------------------------------------------

def extract_resume_text(file_bytes: bytes) -> str:
    """Flatten a .docx master resume into plain text, keeping enough structure
    (headings / bullets marked) for the AI to reason about sections."""
    document = Document(io.BytesIO(file_bytes))
    lines = []
    for p in document.paragraphs:
        text = p.text.strip()
        if not text:
            continue
        if _is_heading(p):
            lines.append(f"\n## {text}")
        elif _has_numpr(p):
            lines.append(f"- {text}")
        else:
            lines.append(text)

    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                lines.append(" | ".join(cells))

    return "\n".join(lines).strip()


# ---------------------------------------------------------------------------
# Resume generation
# ---------------------------------------------------------------------------

def _contact_lines(data: dict) -> list[str]:
    contact = data.get("contact", {}) or {}
    name = (data.get("name") or "").strip()
    title = (data.get("title") or "").strip()
    bits = [contact.get("phone"), contact.get("email"), contact.get("location"),
            contact.get("linkedin"), contact.get("other")]
    contact_line = " | ".join([b.strip() for b in bits if b and b.strip()])
    return [line for line in [name, title, contact_line] if line]


def _render_contact(contact_paras: list, lines: list[str]) -> None:
    if not contact_paras or not lines:
        return
    n, m = len(contact_paras), len(lines)
    if n >= m:
        for i in range(m):
            set_para_text(contact_paras[i], lines[i])
        for i in range(m, n):
            _remove_paragraph(contact_paras[i])
    else:
        for i in range(n - 1):
            set_para_text(contact_paras[i], lines[i])
        set_para_text(contact_paras[-1], " | ".join(lines[n - 1:]))


def _try_render_header_contact(document: Document, lines: list[str]) -> None:
    """Best-effort: some templates put name/contact info in the page header
    instead of the body. If we find short text there, update it too."""
    try:
        header = document.sections[0].header
        header_paras = [p for p in header.paragraphs if p.text.strip()]
        if header_paras and lines:
            _render_contact(header_paras, lines)
    except Exception:
        pass


def _experience_line(job: dict, header_lines: list) -> list[str]:
    title = job.get("job_title", "") or ""
    company = job.get("company", "") or ""
    location = job.get("location", "") or ""
    dates = job.get("dates", "") or ""
    if len(header_lines) >= 2:
        line1 = ", ".join([x for x in [title, company] if x])
        line2 = " | ".join([x for x in [location, dates] if x])
        return [line1, line2] if line2 else [line1]
    else:
        combined = ", ".join([x for x in [title, company] if x])
        tail = " | ".join([x for x in [location, dates] if x])
        return [f"{combined} — {tail}" if tail else combined]


def _render_repeating_section(section: SectionInfo, entries: list, entry_to_lines_and_bullets) -> None:
    """Generic renderer for sections made of repeated entries (jobs, projects,
    education), each with header line(s) + optional bullets."""
    header_lines, bullet_exemplar, generic_exemplar = section.exemplars()
    header_exemplars = header_lines if header_lines else ([generic_exemplar] if generic_exemplar else [])

    new_elements = []
    for entry in entries:
        lines, bullets = entry_to_lines_and_bullets(entry, header_lines)
        for i, line_text in enumerate(lines):
            exemplar = header_exemplars[min(i, len(header_exemplars) - 1)] if header_exemplars else bullet_exemplar
            if exemplar is None:
                continue
            new_elements.append(_clone_paragraph_element(exemplar, line_text))
        for b_text in bullets:
            exemplar = bullet_exemplar or generic_exemplar
            if exemplar is None:
                continue
            prefix = "" if bullet_exemplar else "- "
            new_elements.append(_clone_paragraph_element(exemplar, f"{prefix}{b_text}"))

    for p in section.content_paras:
        _remove_paragraph(p)

    anchor = section.heading_para._p
    for el in new_elements:
        anchor = _insert_after(anchor, el)


def _render_plain_section(section: SectionInfo, text: str) -> None:
    header_lines, bullet_exemplar, generic_exemplar = section.exemplars()
    exemplar = generic_exemplar or bullet_exemplar
    for p in section.content_paras:
        _remove_paragraph(p)
    if exemplar is None or not text:
        return
    el = _clone_paragraph_element(exemplar, text)
    _insert_after(section.heading_para._p, el)


def _render_list_section(section: SectionInfo, items: list[str]) -> None:
    header_lines, bullet_exemplar, generic_exemplar = section.exemplars()
    for p in section.content_paras:
        _remove_paragraph(p)
    if not items:
        return
    anchor = section.heading_para._p
    if bullet_exemplar is not None:
        for item in items:
            el = _clone_paragraph_element(bullet_exemplar, item)
            anchor = _insert_after(anchor, el)
    elif generic_exemplar is not None:
        el = _clone_paragraph_element(generic_exemplar, "  •  ".join(items))
        _insert_after(anchor, el)


def _add_new_section(profile: TemplateProfile, title: str, items: list[str], anchor_element):
    """Create a brand-new heading + bullet list for content the template
    didn't have a slot for (e.g. Certifications found in the master resume
    but missing from the template). Takes and returns a raw lxml element to
    insert after, so multiple new sections can be chained onto the true end
    of the document regardless of which original section came last."""
    heading_exemplar = profile.fallback_heading
    bullet_exemplar = profile.fallback_bullet
    plain_exemplar = profile.fallback_plain
    if heading_exemplar is None or anchor_element is None:
        return anchor_element

    anchor = anchor_element
    heading_el = _clone_paragraph_element(heading_exemplar, title.upper())
    anchor = _insert_after(anchor, heading_el)

    for item in items:
        exemplar = bullet_exemplar or plain_exemplar
        if exemplar is None:
            continue
        prefix = "" if bullet_exemplar else "- "
        el = _clone_paragraph_element(exemplar, f"{prefix}{item}")
        anchor = _insert_after(anchor, el)

    return anchor


def generate_tailored_docx(template_bytes: bytes, resume_data: dict) -> bytes:
    """Populate the uploaded template with tailored resume_data and return the
    resulting .docx as bytes."""
    document = Document(io.BytesIO(template_bytes))
    profile = _build_profile(document)

    lines = _contact_lines(resume_data)
    _render_contact(profile.contact_paras, lines)
    _try_render_header_contact(document, lines)

    summary_section = profile.get("summary")
    if summary_section is not None:
        _render_plain_section(summary_section, resume_data.get("summary", "") or "")

    skills_section = profile.get("skills")
    if skills_section is not None:
        _render_list_section(skills_section, resume_data.get("skills", []) or [])

    experience_section = profile.get("experience")
    if experience_section is not None:
        def exp_mapper(job, header_lines):
            lines_ = _experience_line(job, header_lines)
            bullets_ = job.get("bullets", []) or []
            return lines_, bullets_
        _render_repeating_section(experience_section, resume_data.get("experience", []) or [], exp_mapper)

    education_section = profile.get("education")
    if education_section is not None:
        _, edu_bullet_exemplar, _ = education_section.exemplars()
        has_edu_bullets = edu_bullet_exemplar is not None

        def edu_mapper(edu, header_lines):
            degree = edu.get("degree", "") or ""
            school = edu.get("school", "") or ""
            location = edu.get("location", "") or ""
            dates = edu.get("dates", "") or ""
            details = edu.get("details", "") or ""
            if len(header_lines) >= 2:
                line1 = ", ".join([x for x in [degree, school] if x])
                line2 = " | ".join([x for x in [location, dates] if x])
                out_lines = [line1] + ([line2] if line2 else [])
                bullets_ = [details] if (details and has_edu_bullets) else []
                if details and not has_edu_bullets:
                    out_lines[-1] = f"{out_lines[-1]} ({details})" if out_lines[-1] else details
            else:
                combined = ", ".join([x for x in [degree, school] if x])
                tail = " | ".join([x for x in [location, dates] if x])
                line = f"{combined} — {tail}" if tail else combined
                if details and not has_edu_bullets:
                    line = f"{line} ({details})" if line else details
                out_lines = [line] if line else []
                bullets_ = [details] if (details and has_edu_bullets) else []
            return out_lines, bullets_
        _render_repeating_section(education_section, resume_data.get("education", []) or [], edu_mapper)

    certifications_section = profile.get("certifications")
    cert_items = resume_data.get("certifications", []) or []
    if certifications_section is not None:
        _render_list_section(certifications_section, cert_items)
        cert_items = []  # already placed

    training_section = profile.get("training")
    training_items = resume_data.get("training", []) or []
    if training_section is not None:
        _render_list_section(training_section, training_items)
        training_items = []

    languages_section = profile.get("languages")
    language_items = resume_data.get("languages", []) or []
    if languages_section is not None:
        _render_list_section(languages_section, language_items)
        language_items = []

    projects_section = profile.get("projects")
    project_entries = resume_data.get("projects", []) or []
    if projects_section is not None:
        def proj_mapper(proj, header_lines):
            name = proj.get("name", "") or ""
            desc = proj.get("description", "") or ""
            line = name if not desc else f"{name} — {desc}" if len(header_lines) < 2 else name
            out_lines = [line] if line else []
            if len(header_lines) >= 2 and desc:
                out_lines.append(desc)
            bullets_ = proj.get("bullets", []) or []
            return out_lines, bullets_
        _render_repeating_section(projects_section, project_entries, proj_mapper)
        project_entries = []

    # Anything the template had no slot for gets appended as a brand-new
    # section at the true end of the document body (just before the section
    # properties element), regardless of which original heading came last.
    from docx.oxml.ns import qn

    body = document.element.body
    children = list(body)
    if children and children[-1].tag == qn("w:sectPr"):
        anchor_el = children[-2] if len(children) >= 2 else None
    else:
        anchor_el = children[-1] if children else None

    def _flatten_experience(jobs):
        out = []
        for job in jobs:
            head = ", ".join([x for x in [job.get("job_title", ""), job.get("company", "")] if x])
            tail = " | ".join([x for x in [job.get("location", ""), job.get("dates", "")] if x])
            out.append(f"{head} — {tail}" if tail else head)
            out.extend(job.get("bullets", []) or [])
        return [x for x in out if x]

    def _flatten_education(edus):
        out = []
        for edu in edus:
            head = ", ".join([x for x in [edu.get("degree", ""), edu.get("school", "")] if x])
            tail = " | ".join([x for x in [edu.get("location", ""), edu.get("dates", "")] if x])
            out.append(f"{head} — {tail}" if tail else head)
            if edu.get("details"):
                out.append(edu["details"])
        return [x for x in out if x]

    def _flatten_projects(projs):
        out = []
        for proj in projs:
            name = proj.get("name", "") or ""
            desc = proj.get("description", "") or ""
            out.append(f"{name} — {desc}" if desc else name)
            out.extend(proj.get("bullets", []) or [])
        return [x for x in out if x]

    missing_sections = []
    if summary_section is None and resume_data.get("summary"):
        missing_sections.append(("Summary", [resume_data["summary"]]))
    if skills_section is None and resume_data.get("skills"):
        missing_sections.append(("Skills", resume_data.get("skills", [])))
    if experience_section is None and resume_data.get("experience"):
        missing_sections.append(("Experience", _flatten_experience(resume_data.get("experience", []))))
    if education_section is None and resume_data.get("education"):
        missing_sections.append(("Education", _flatten_education(resume_data.get("education", []))))
    if certifications_section is None and cert_items:
        missing_sections.append(("Certifications", cert_items))
    if training_section is None and training_items:
        missing_sections.append(("Professional Training", training_items))
    if languages_section is None and language_items:
        missing_sections.append(("Languages", language_items))
    if projects_section is None and project_entries:
        missing_sections.append(("Projects", _flatten_projects(project_entries)))

    additional = resume_data.get("additional_sections", {}) or {}
    for title, items in additional.items():
        if items:
            missing_sections.append((title, items))

    for title, items in missing_sections:
        if not items or anchor_el is None:
            continue
        anchor_el = _add_new_section(profile, title, items, anchor_el)

    out = io.BytesIO()
    document.save(out)
    return out.getvalue()
